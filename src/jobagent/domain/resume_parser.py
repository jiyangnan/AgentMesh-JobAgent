"""Resume text extraction — PDF / DOCX / TXT / Markdown."""

from __future__ import annotations

import re
from pathlib import Path
from typing import Any


class ResumeParser:
    """Extract plain text from resume files in various formats."""

    SUPPORTED_EXTS = {".pdf", ".docx", ".txt", ".md", ".markdown"}

    def parse(self, file_path: str | Path) -> str:
        """Extract and clean text from a resume file.

        Returns:
            Cleaned text string. Raises ValueError for unsupported formats.
        """
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"Resume not found: {path}")

        ext = path.suffix.lower()
        if ext == ".pdf":
            raw = self._extract_pdf(path)
        elif ext == ".docx":
            raw = self._extract_docx(path)
        elif ext in (".txt", ".md", ".markdown"):
            raw = path.read_text(encoding="utf-8")
        else:
            raise ValueError(
                f"Unsupported resume format: {ext}. "
                f"Supported: {', '.join(self.SUPPORTED_EXTS)}"
            )

        cleaned = self._clean_text(raw)
        if len(cleaned) < 100:
            raise ValueError(
                "Extracted text is too short (<100 chars). "
                "The file may be a scanned image without OCR text layer."
            )
        return cleaned

    # ── Format-specific extractors ──────────────────────────

    def _extract_pdf(self, path: Path) -> str:
        """Extract text from PDF using PyMuPDF."""
        import fitz  # pymupdf

        doc = fitz.open(path)
        parts: list[str] = []
        for page in doc:
            text = page.get_text()
            if text.strip():
                parts.append(text)
        doc.close()
        return "\n".join(parts)

    def _extract_docx(self, path: Path) -> str:
        """Extract text from DOCX using python-docx."""
        from docx import Document

        doc = Document(path)
        parts: list[str] = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text)
        return "\n".join(parts)

    # ── Text cleaning ───────────────────────────────────────

    def _clean_text(self, text: str) -> str:
        """Normalize and clean extracted resume text."""
        # Normalize line endings
        text = text.replace("\r\n", "\n").replace("\r", "\n")
        # Collapse multiple blank lines
        text = re.sub(r"\n{3,}", "\n\n", text)
        # Remove standalone page numbers (heuristic)
        text = re.sub(r"\n\s*\d+\s*\n", "\n", text)
        # Remove common header/footer artifacts
        text = re.sub(r"\n\s*第\s*\d+\s*页\s*\n", "\n", text)
        text = re.sub(r"\n\s*Page\s*\d+\s*\n", "\n", text, flags=re.IGNORECASE)
        # Strip leading/trailing whitespace
        text = text.strip()
        # Truncate to ~8000 chars for LLM context window (rough estimate)
        if len(text) > 8000:
            text = text[:8000] + "\n...[truncated for LLM context window]"
        return text


def guess_encoding(file_path: str | Path) -> str:
    """Best-effort text file encoding detection."""
    path = Path(file_path)
    for encoding in ("utf-8", "gbk", "gb2312", "latin-1"):
        try:
            path.read_text(encoding=encoding)
            return encoding
        except UnicodeDecodeError:
            continue
    return "utf-8"
